

class CV:
    def __init__(self):
        self.BI = ""
        self.SN = ""
        self.Isim = ""
        self.DogumTarihi = ""
        self.Mail = ""
        self.Deneyimler = ""
        self.Egitim = ""
        self.Sertifikalar = ""
        self.Diller = ""
        self.Organizasyonlar = ""
        self.Projeler = ""
        self.Oduller = ""
        self.Yayinlar = ""
        self.Referanslar = ""
        self.Vizyon = ""

    def CvFill(self):
        if self.SN == "S":
            self.SN = 'List Number'
        else:
            self.SN = 'List Bullet'

        from docx import Document

        document = Document()

        document.add_heading('{isim} - Özgeçmiş'.format(isim = self.Isim), 0)
# TODO: buraya bold-italic ekle
        document.add_heading('Kişisel Bilgiler: ', level=1)
        document.add_paragraph('Ad-Soyad: {isim}'.format(isim=self.Isim))
        document.add_paragraph('Doğum Tarihi: {tarih}'.format(tarih=self.DogumTarihi))
        document.add_paragraph('Mail: {isim}'.format(isim=self.Mail))

        document.add_heading('Deneyimler: ', level=1)
        document.add_paragraph('{deneyim}'.format(deneyim=self.Deneyimler), style=self.SN)

        document.add_heading('Eğitim: ', level=1)
        document.add_paragraph('{egitim}'.format(egitim=self.Egitim), style=self.SN)

        document.add_heading('Sertifikalar: ', level=1)
        document.add_paragraph('{sertifika}'.format(sertifika=self.Sertifikalar), style=self.SN)

        document.add_heading('Diller: ', level=1)
        document.add_paragraph('{dil}'.format(dil=self.Diller), style=self.SN)

        document.add_heading('Organizasyonlar: ', level=1)
        document.add_paragraph('{organizasyon}'.format(organizasyon=self.Organizasyonlar), style=self.SN)

        document.add_heading('Projeler: ', level=1)
        document.add_paragraph('{proje}'.format(proje=self.Projeler), style=self.SN)

        document.add_heading('Ödüller: ', level=1)
        document.add_paragraph('{odul}'.format(odul=self.Oduller), style=self.SN)

        document.add_heading('Yayınlar: ', level=1)
        document.add_paragraph('{yayin}'.format(yayin=self.Yayinlar), style=self.SN)

        document.add_heading('Vizyon: ', level=1)
        v = document.add_paragraph("*")
        if self.BI == "B":
            v.add_run(self.Vizyon).bold = True
        elif self.BI == "I":
            v.add_run(self.Vizyon).italic = True
        else:
            v.add_run(self.Vizyon)

        document.add_heading('Referanslar: ', level=1)
        document.add_paragraph('{referans}'.format(referans=self.Referanslar), style=self.SN)

        document.save('Özgeçmiş.docx')
CV = CV()